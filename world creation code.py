import pandas as pd

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# ==================================================
#read the data
# ==================================================

df = pd.read_csv(r"F:\Career development program\Practice\Agricultural project\spss-e\crop1.csv")

# convert numerical data
df['G09'] = pd.to_numeric(df['G09'], errors='coerce')
df['G10'] = pd.to_numeric(df['G10'], errors='coerce')

# delet the empty rows
df = df.dropna(subset=['A02', 'G08', 'G09', 'G10'])

# ==================================================
# create report
# ==================================================

doc = Document()

# ==================================================
# the main title
# ==================================================

doc.add_heading(
    'Agricultural Data Analysis Report',
    level=1
)

# ==================================================
# introduction
# ==================================================

p = doc.add_paragraph()

run = p.add_run(
    'This report presents an exploratory '
    'analysis of agricultural census data '
    'including governorate efficiency, '
    'irrigation systems, and production patterns.'
)

run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# ==================================================
#governorates analysis
# ==================================================

doc.add_heading(
    '1. Governorate Analysis',
    level=2
)

gov_analysis = df.groupby('A02').agg({
    'G09':'sum',
    'G10':'sum'
})

gov_analysis.columns = [
    'Total Area',
    'Total Production'
]

gov_analysis['Efficiency'] = (
    gov_analysis['Total Production'] /
    gov_analysis['Total Area']
)

gov_analysis = gov_analysis.sort_values(
    'Efficiency',
    ascending=False
)

# governorates table
table1 = doc.add_table(
    rows=1,
    cols=4
)

table1.style = 'Table Grid'

hdr = table1.rows[0].cells

hdr[0].text = 'Governorate'
hdr[1].text = 'Total Area'
hdr[2].text = 'Total Production'
hdr[3].text = 'Efficiency'

for idx, row in gov_analysis.head(10).iterrows():

    cells = table1.add_row().cells

    cells[0].text = str(idx)

    cells[1].text = str(
        round(row['Total Area'], 2)
    )

    cells[2].text = str(
        round(row['Total Production'], 2)
    )

    cells[3].text = str(
        round(row['Efficiency'], 2)
    )

# explain
p = doc.add_paragraph()

run = p.add_run(
    'Write your interpretation here about '
    'the differences in agricultural efficiency '
    'between governorates.'
)

run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# ==================================================
# irrigation analysis
# ==================================================

doc.add_heading(
    '2. Irrigation Analysis',
    level=2
)

irrigation_analysis = df.groupby('G08').agg({
    'G09':'sum',
    'G10':'sum'
})

irrigation_analysis.columns = [
    'Total Area',
    'Total Production'
]

irrigation_analysis['Efficiency'] = (
    irrigation_analysis['Total Production'] /
    irrigation_analysis['Total Area']
)

irrigation_analysis = irrigation_analysis.sort_values(
    'Efficiency',
    ascending=False
)

# irrigation table
table2 = doc.add_table(
    rows=1,
    cols=4
)

table2.style = 'Table Grid'

hdr = table2.rows[0].cells

hdr[0].text = 'Irrigation Type'
hdr[1].text = 'Total Area'
hdr[2].text = 'Total Production'
hdr[3].text = 'Efficiency'

for idx, row in irrigation_analysis.iterrows():

    cells = table2.add_row().cells

    cells[0].text = str(idx)

    cells[1].text = str(
        round(row['Total Area'], 2)
    )

    cells[2].text = str(
        round(row['Total Production'], 2)
    )

    cells[3].text = str(
        round(row['Efficiency'], 2)
    )

# explaination traitment
p = doc.add_paragraph()

run = p.add_run(
    'Write your interpretation here about '
    'the impact of irrigation systems on '
    'agricultural production efficiency.'
)

run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# ==================================================
# to 10 efficiency
# ==================================================

doc.add_heading(
    '3. Top 10 Agricultural Efficiency Results',
    level=2
)

gov_irrigation = df.groupby(
    ['A02', 'G08']
).agg({
    'G09':'sum',
    'G10':'sum'
})

gov_irrigation.columns = [
    'Total Area',
    'Total Production'
]

gov_irrigation['Efficiency'] = (
    gov_irrigation['Total Production'] /
    gov_irrigation['Total Area']
)

gov_irrigation = gov_irrigation.sort_values(
    'Efficiency',
    ascending=False
)

top10 = gov_irrigation.head(10)

# top 10 efficiency table
table3 = doc.add_table(
    rows=1,
    cols=5
)

table3.style = 'Table Grid'

hdr = table3.rows[0].cells

hdr[0].text = 'Governorate'
hdr[1].text = 'Irrigation'
hdr[2].text = 'Total Area'
hdr[3].text = 'Total Production'
hdr[4].text = 'Efficiency'

for idx, row in top10.iterrows():

    cells = table3.add_row().cells

    cells[0].text = str(idx[0])

    cells[1].text = str(idx[1])

    cells[2].text = str(
        round(row['Total Area'], 2)
    )

    cells[3].text = str(
        round(row['Total Production'], 2)
    )

    cells[4].text = str(
        round(row['Efficiency'], 2)
    )

# explaination 
p = doc.add_paragraph()

run = p.add_run(
    'Write your interpretation here about '
    'the highest agricultural efficiency '
    'combinations between governorates and '
    'irrigation systems.'
)

run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# ==================================================
# save the report
# ==================================================

doc.save(
    'Agricultural_Analysis_Report.docx'
)

print('Report created successfully.')
